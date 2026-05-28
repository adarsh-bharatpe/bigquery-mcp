#!/usr/bin/env python3
"""Export upi-retention-results.html to Word (.docx) without dropping slide content."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

HTML_PATH = Path(__file__).resolve().parent / "upi-retention-results.html"
OUT_PATH = Path(__file__).resolve().parent / "upi-retention-results.docx"

HANDLED_TAGS = frozenset(
    {
        "h1",
        "h2",
        "h3",
        "h4",
        "p",
        "table",
        "footer",
        "li",
        "button",
        "script",
        "style",
        "svg",
        "img",
    }
)


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def element_text(el) -> str:
    return clean_text(el.get_text(separator=" ", strip=True))


def classes(el) -> list[str]:
    return el.get("class") or []


def add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int | None = None):
    text = clean_text(text)
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)


def add_table_from_element(doc: Document, table_el) -> None:
    rows = table_el.find_all("tr")
    if not rows:
        return
    max_cols = max(len(r.find_all(["th", "td"])) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=max_cols)
    tbl.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for j, cell in enumerate(cells):
            tbl.rows[i].cells[j].text = element_text(cell)
            if cell.name == "th":
                for p in tbl.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True


def is_inside_handled_special(el) -> bool:
    """Skip li inside cards we handle as a unit."""
    parent = el.parent
    while parent:
        pc = parent.get("class") or []
        if any(
            x in pc
            for x in (
                "pattern-card",
                "action-box",
                "comp-app-block",
                "comp-app-head",
                "comp-strategy-wrap",
                "reddit-preview",
                "metric-card",
                "lens-btn",
                "emerging-patterns",
                "comp-apps-grid",
            )
        ):
            return True
        parent = parent.parent
    return False


def emit_element(doc: Document, el) -> None:
    c = classes(el)

    if "metric-card" in c:
        lbl, num, sub = el.select_one(".lbl"), el.select_one(".num"), el.select_one(".sub")
        bits = [element_text(x) for x in (lbl, num) if x]
        if sub:
            bits.append(f"({element_text(sub)})")
        add_paragraph(doc, " · ".join(bits))
        return

    if "lens-btn" in c:
        tag, hint = el.select_one(".tag"), el.select_one(".hint")
        label = element_text(el)
        if tag:
            label = label.replace(element_text(tag), "", 1)
        if hint:
            label = label.replace(element_text(hint), "", 1)
        add_paragraph(doc, f"{element_text(tag)} {clean_text(label)} {element_text(hint)}".strip(), italic=True)
        return

    if "action-box" in c:
        src = el.select_one(".action-source")
        if src:
            add_paragraph(doc, element_text(src), italic=True, size=9)
        h = el.find(["h3", "h4"], recursive=False)
        if h:
            doc.add_heading(element_text(h), level=4)
        for li in el.select("ul > li"):
            doc.add_paragraph(element_text(li), style="List Bullet")
        return

    if "pattern-card" in c:
        num = el.select_one(".num")
        h = el.find(["h4", "h3"], recursive=False)
        title = element_text(h)
        if num:
            title = f"{element_text(num)}. {title}"
        doc.add_heading(title, level=4)
        for li in el.select("ul > li"):
            doc.add_paragraph(element_text(li), style="List Bullet")
        note = el.select_one(".note")
        if note:
            add_paragraph(doc, element_text(note), italic=True)
        return

    if "reddit-preview" in c:
        href = el.get("href", "")
        bits = [
            element_text(el.select_one(".reddit-preview-sub")),
            element_text(el.select_one(".reddit-preview-title")),
            f'"{element_text(el.select_one(".reddit-preview-quote"))}"' if el.select_one(".reddit-preview-quote") else "",
            href,
        ]
        add_paragraph(doc, " — ".join(b for b in bits if b))
        return

    if "comp-app-block" in c:
        h = el.find("h3", recursive=False)
        if h:
            doc.add_heading(element_text(h), level=3)
        for table in el.find_all("table"):
            add_table_from_element(doc, table)
        return

    if "comp-extra-item" in c:
        add_paragraph(doc, element_text(el))
        return

    if el.name == "table":
        add_table_from_element(doc, el)
        return

    if el.name == "h1":
        doc.add_heading(element_text(el), level=1)
        return

    if el.name == "h2":
        doc.add_heading(element_text(el), level=2)
        return

    if el.name == "h3":
        doc.add_heading(element_text(el), level=3)
        return

    if el.name == "h4":
        doc.add_heading(element_text(el), level=4)
        return

    if el.name == "footer" and "takeaway" in c:
        add_paragraph(doc, element_text(el), italic=True)
        return

    if el.name == "p":
        add_paragraph(doc, element_text(el))
        return

    if el.name == "li" and not is_inside_handled_special(el):
        doc.add_paragraph(element_text(el), style="List Bullet")
        return

    if any(x in c for x in ("title-chip", "title-part", "trail-pill", "card", "pref-pill", "sat-label", "row-label", "values-label")):
        if "pref-grid" in c:
            pills = el.select(".pref-pill")
            if pills:
                add_paragraph(doc, " · ".join(element_text(p) for p in pills))
            return
        add_paragraph(doc, element_text(el))
        return


def should_emit(el) -> bool:
    if isinstance(el, NavigableString):
        return False
    if not getattr(el, "name", None):
        return False
    if el.name == "button":
        return False
    c = classes(el)
    if any(
        x in c
        for x in (
            "metric-card",
            "lens-btn",
            "action-box",
            "pattern-card",
            "reddit-preview",
            "comp-app-block",
            "comp-extra-item",
            "title-chip",
            "title-part",
            "trail-pill",
            "sat-label",
            "row-label",
            "values-label",
            "pref-pill",
        )
    ):
        return True
    if el.name in ("h1", "h2", "h4", "table", "footer"):
        return True
    if el.name == "h3" and not is_inside_handled_special(el):
        return True
    if el.name == "p":
        return True
    if el.name == "li" and not is_inside_handled_special(el):
        parent = el.parent
        if parent and parent.name in ("ul", "ol"):
            return True
    if "card" in c and el.parent and "agenda-cards" in (el.parent.get("class") or []):
        return True
    return False


def extract_slide_comments(slide_html: str) -> list[str]:
    out = []
    for m in re.finditer(r"<!--(.*?)-->", slide_html, re.DOTALL):
        body = clean_text(m.group(1))
        if not body or re.match(r"^slide\b", body, re.I):
            continue
        if body.startswith("<"):
            out.append(body)
        elif len(body) > 2:
            out.append(body)
    return out


def export(html_path: Path, out_path: Path) -> None:
    raw = html_path.read_text(encoding="utf-8")
    soup = BeautifulSoup(raw, "lxml")

    doc = Document()
    doc.sections[0].left_margin = Inches(0.75)
    doc.sections[0].right_margin = Inches(0.75)

    doc.add_heading("BharatPe UPI — Engagement & Retention Analysis", level=0)
    add_paragraph(
        doc,
        "Full export from upi-retention-results.html — all slides, tables, takeaways, competitor strategies, and community quotes.",
        italic=True,
    )

    slides = soup.select("section.slide[data-slide]:not(.slide-omitted)")
    slides.sort(key=lambda s: int(s.get("data-slide", 0)))

    for slide in slides:
        num = slide.get("data-slide", "?")
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        doc.add_heading(f"Slide {num}", level=1)

        frame = slide.select_one(".slide-frame") or slide
        emitted_ids: set[int] = set()

        for el in frame.find_all(True):
            if id(el) in emitted_ids:
                continue
            if not should_emit(el):
                continue
            # Skip nested elements if ancestor already emitted as special block
            skip = False
            for parent in el.parents:
                if id(parent) in emitted_ids and should_emit(parent):
                    skip = True
                    break
            if skip:
                continue
            emit_element(doc, el)
            emitted_ids.add(id(el))

        for comment in extract_slide_comments(str(slide)):
            doc.add_heading("HTML comment (not shown in live deck)", level=3)
            add_paragraph(doc, comment, italic=True)

    doc.save(str(out_path))
    print(f"Wrote {out_path} ({len(slides)} slides)")


if __name__ == "__main__":
    html = Path(sys.argv[1]) if len(sys.argv) > 1 else HTML_PATH
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else OUT_PATH
    export(html, out)
